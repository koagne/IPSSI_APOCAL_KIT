from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_styled_table(doc, headers, rows, header_color="1F4E78"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, header_color)
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            table.rows[i + 1].cells[j].text = str(value)
    return table

def main():
    doc = Document()

    # Titre
    title = doc.add_heading('APOCAL\'IPSSI 2026 - ÉQUIPE 2', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Product Vision Board - EduTutor IA')
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()

    # Identification
    doc.add_heading('IDENTIFICATION DU DOCUMENT', level=1)
    add_styled_table(doc, ['Champ', 'Valeur'], [
        ['Équipe n°', '2'],
        ['Membres', 'Danielle Jamila KOAGNE NGANKAM, Krishmini KULAKRISHNA, Wicramachine SERGIO, Houda OUADAH, Adja Fatou SAGNA, Mohammed DERKAOUI, Ousmane NDIAYE'],
        ['Sprint concerné', 'Cadrage'],
        ['Version', 'v1.0 (initiale)'],
        ['Date de remise', datetime.datetime.now().strftime('%d/%m/%Y 12h50')],
        ['Statut', 'Draft'],
    ])
    doc.add_paragraph()

    # 1. VISION
    doc.add_heading('1. Vision', level=1)
    vision_para = doc.add_paragraph()
    vision_run = vision_para.add_run(
        '« Permettre à chaque étudiant·e du supérieur de transformer n\'importe quel cours '
        'en quiz de révision personnalisé en moins de 5 minutes, et offrir aux enseignant·e·s '
        'un gain de 10h/mois de préparation de supports d\'évaluation, le tout sans qu\'aucune '
        'donnée pédagogique ne quitte le territoire européen. »'
    )
    vision_run.italic = True
    doc.add_paragraph()

    # 2. TARGET GROUP
    doc.add_heading('2. Target Group (cibles utilisateurs)', level=1)

    doc.add_heading('2.1. Cible primaire - Étudiant·e du supérieur', level=2)
    add_styled_table(doc, ['Champ', 'Valeur'], [
        ['Profil', '18-25 ans · L1 à M2 · usage quotidien smartphone + laptop'],
        ['Volume FR', '~2,7M d\'étudiants dans le supérieur (MESR 2024)'],
        ['Pain point', '5h à 15h/semaine perdues à chercher ou créer des supports de révision'],
        ['Critère clé', 'Gratuité/freemium + confidentialité des cours + rapidité (< 5 min)'],
    ])
    doc.add_paragraph()

    doc.add_heading('2.2. Cible secondaire - Enseignant·e (Mme Lefèvre)', level=2)
    add_styled_table(doc, ['Champ', 'Valeur'], [
        ['Profil', '35-55 ans · lycée sous contrat / BTS / supérieur'],
        ['Volume FR', '~770 000 enseignants (Éducation nationale 2024)'],
        ['Pain point', '~12h/mois en correction et préparation de supports d\'évaluation'],
        ['Critère clé', 'Interface ultra-simple + conformité RGPD + export Word/PDF'],
    ])
    doc.add_paragraph()

    doc.add_heading('2.3. Cible tertiaire - Établissement scolaire (B2B)', level=2)
    add_styled_table(doc, ['Champ', 'Valeur'], [
        ['Profil', 'Direction de lycée privé · ENT scolaire · responsable pédagogique'],
        ['Volume FR', '~7 500 lycées + ~3 500 établissements supérieurs'],
        ['Pain point', 'Budget edtech contraint (< 5€/élève/an) + obligation RGPD'],
        ['Critère clé', 'Hébergement UE souverain + facturation prévisible'],
    ])
    doc.add_paragraph()

    # 3. NEEDS
    doc.add_heading('3. Needs (besoins résolus)', level=1)
    doc.add_heading('3.1. Besoins étudiant', level=2)
    for need in [
        'Générer en moins de 5 min un quiz de 10 QCM à partir d\'un PDF ou texte',
        'Identifier ses lacunes chapitre par chapitre avant un examen',
        'Retrouver et rejouer ses anciens quiz via un historique persisté',
        'Réviser sur mobile comme sur laptop (interface responsive)',
    ]:
        doc.add_paragraph(need, style='List Bullet')
    doc.add_paragraph()

    doc.add_heading('3.2. Besoins enseignant (Mme Lefèvre)', level=2)
    for need in [
        'Préparer un support d\'évaluation (10 QCM + corrigé) en moins de 3 minutes',
        'Disposer de questions pédagogiques ancrées dans le cours fourni',
        'Exporter le quiz en format imprimable (PDF/Word) pour distribution en classe',
        'Suivre la performance globale de la classe via un tableau de bord',
    ]:
        doc.add_paragraph(need, style='List Bullet')
    doc.add_paragraph()

    doc.add_heading('3.3. Besoins établissement', level=2)
    for need in [
        'Disposer d\'un outil 100% RGPD compliant, sans transfert hors UE',
        'Bénéficier d\'une tarification prévisible par élève/an',
        'Obtenir une traçabilité complète des sources pédagogiques (RAG)',
    ]:
        doc.add_paragraph(need, style='List Bullet')
    doc.add_paragraph()

    # 4. PRODUCT
    doc.add_heading('4. Product (signature produit)', level=1)
    doc.add_heading('4.1. Caractéristiques signature', level=2)
    for feat in [
        'Génération de 10 QCM en moins de 60 secondes à partir d\'un cours',
        'Fonctionnement 100% local via Ollama (Llama 3.1 8B)',
        'Interface mobile-first, responsive, accessible sans installation',
        'Questions pédagogiques traçables à leur source (RAG en Release 2)',
    ]:
        doc.add_paragraph(feat, style='List Bullet')
    doc.add_paragraph()

    doc.add_heading('4.2. MVP must-have (Release 1)', level=2)
    for feat in [
        'F1 : Inscription et connexion par email (validation + reset + profil)',
        'F2 : Saisie d\'un cours (upload PDF ≤ 5 Mo ou texte ≥ 200 caractères)',
        'F3 : Génération automatique d\'un quiz de 10 QCM via LLM local',
        'F4 : Soumission et correction automatique',
        'F5 : Affichage du score /10 + détail bonnes/mauvaises réponses',
        'F6 : Historique persisté des quiz par utilisateur',
    ]:
        doc.add_paragraph(feat, style='List Bullet')
    doc.add_paragraph()

    doc.add_heading('4.3. Pistes Release 2 candidates', level=2)
    for title, justification in [
        ('Export PDF/Word du quiz + corrigé', 'Répond au besoin de Mme Lefèvre'),
        ('RAG avec traçabilité source', 'Renforce la pédagogie ancrée'),
        ('Tableau de bord de progression', 'Déjà codé en démo, à enrichir'),
    ]:
        p = doc.add_paragraph(style='List Bullet')
        run_title = p.add_run(title + ' — ')
        run_title.bold = True
        p.add_run(justification)
    doc.add_paragraph()

    # 5. BUSINESS GOALS
    doc.add_heading('5. Business Goals', level=1)
    doc.add_heading('5.1. Objectifs d\'adoption', level=2)
    for goal in [
        '1 000 étudiants actifs hebdomadaires (WAU) d\'ici T+6 mois',
        'Taux de rétention J7 : ≥ 40%',
        '500 quiz générés/jour d\'ici T+4 mois',
    ]:
        doc.add_paragraph(goal, style='List Bullet')
    doc.add_paragraph()

    doc.add_heading('5.2. Objectifs de satisfaction', level=2)
    for goal in [
        'NPS > 30 d\'ici T+9 mois',
        '< 5% de quiz signalés "erreur factuelle"',
        'Temps moyen de génération : < 60 secondes (P95)',
    ]:
        doc.add_paragraph(goal, style='List Bullet')
    doc.add_paragraph()

    doc.add_heading('5.3. Objectifs business', level=2)
    for goal in [
        '10 établissements B2B dans les 12 mois',
        'CAC < 8€ par utilisateur converti',
        'Modèle : freemium étudiant + licence B2B (3€/élève/an)',
    ]:
        doc.add_paragraph(goal, style='List Bullet')
    doc.add_paragraph()

    # 6. DIFFÉRENCIATEURS
    doc.add_heading('6. Différenciateurs vs concurrents', level=1)
    add_styled_table(doc, ['Concurrent', 'Positionnement', 'Limite'], [
        ['Wilgo.ai', 'Compagnon IA français étudiants', 'Cloud, données hors UE'],
        ['Leo', 'Tuteur IA Bac/sup français', 'Pas d\'angle enseignant'],
        ['Quizlet AI', 'Cartes mémoire et quiz IA US', 'Focus marché US'],
        ['Khanmigo', 'Tuteur IA Khan Academy', 'US-first, RGPD floue'],
    ])
    doc.add_paragraph()

    doc.add_heading('6.2. Nos 3 différenciateurs', level=2)
    doc.add_paragraph('1. Prompts métier enseignants (enseignant-first)', style='List Number')
    doc.add_paragraph('2. Pédagogie ancrée (RAG sur cours fourni)', style='List Number')
    doc.add_paragraph('3. RGPD local-first (souveraineté des données)', style='List Number')

    # Sauvegarde
    filename = 'equipe-02-product-vision-board-v1.0.docx'
    doc.save(filename)
    print(f'✅ Fichier généré : {filename}')

if __name__ == '__main__':
    main()
