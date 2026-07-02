#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generate.py — Génère les .docx du dossier d'avant-projet EduTutor IA.

Convertit chaque fichier Markdown de sources/*.md en un document Word soigné
(docx/*.docx) : page de garde, styles de titres, tableaux, et insertion des
diagrammes (PNG de diagrammes/img/) aux emplacements marqués
« [[DIAGRAMME: <fichier>.svg]] ».

Dépendance : python-docx  (pip install python-docx)
Usage      : python generate.py
Puis       : ouvrez les .docx et « Exporter en PDF » (Word / LibreOffice).

[Note pédagogique] On garde les SOURCES en Markdown (versionnables, faciles à
diffuser) et on les rend en .docx via ce script : la « source de vérité » reste
le texte, pas le binaire.
"""
import os
import re
import glob

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__))
SOURCES = os.path.join(HERE, "sources")
DOCX_OUT = os.path.join(HERE, "docx")
IMG = os.path.join(HERE, "diagrammes", "img")

INDIGO = RGBColor(0x4F, 0x46, 0xE5)
SLATE = RGBColor(0x33, 0x41, 0x55)

INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`.+?`)")
DIAGRAM_RE = re.compile(r"^\[\[DIAGRAMME:\s*(.+?)\s*\]\]$")


def add_runs(paragraph, text):
    """Ajoute du texte avec gestion du gras (**), italique (*) et code (`)."""
    for token in INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            r = paragraph.add_run(token[2:-2]); r.bold = True
        elif token.startswith("*") and token.endswith("*"):
            r = paragraph.add_run(token[1:-1]); r.italic = True
        elif token.startswith("`") and token.endswith("`"):
            r = paragraph.add_run(token[1:-1]); r.font.name = "Consolas"
        else:
            paragraph.add_run(token)


def cover_page(doc, title):
    """Page de garde."""
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("EduTutor IA"); r.bold = True; r.font.size = Pt(34); r.font.color.rgb = INDIGO
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Dossier d'avant-projet"); r.font.size = Pt(16); r.font.color.rgb = SLATE
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); r.bold = True; r.font.size = Pt(22)
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Document hérité de la première équipe — APOCAL'IPSSI 2026")
    r.italic = True; r.font.color.rgb = SLATE
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Version 1.0 · à compléter par votre équipe"); r.italic = True; r.font.size = Pt(10)
    doc.add_page_break()


def insert_diagram(doc, svg_name):
    """Insère le PNG correspondant (ou un placeholder visible si absent)."""
    base = os.path.splitext(svg_name)[0]
    png = os.path.join(IMG, base + ".png")
    if os.path.isfile(png):
        doc.add_picture(png, width=Cm(16))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"[ DIAGRAMME À INSÉRER : {svg_name} — voir diagrammes/img/ ]")
        r.bold = True; r.font.color.rgb = INDIGO


def flush_table(doc, rows):
    """Crée un tableau Word à partir des lignes Markdown (| .. | .. |)."""
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    # ligne 1 = entête, ligne 2 = séparateur (---), reste = données
    header = cells[0]
    body = cells[2:] if len(cells) > 2 else []
    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        table.style = "Table Grid"
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(h.replace("**", "")); run.bold = True
    for row in body:
        tr = table.add_row().cells
        for i, val in enumerate(row):
            if i < len(tr):
                tr[i].paragraphs[0].text = ""
                add_runs(tr[i].paragraphs[0], val)
    doc.add_paragraph()


def render(md_path, docx_path):
    text = open(md_path, encoding="utf-8").read()
    lines = text.splitlines()
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # Titre H1 -> page de garde
    title = "Document"
    for l in lines:
        if l.startswith("# "):
            title = l[2:].strip(); break
    cover_page(doc, title)

    table_buf = []
    i = 0
    first_h1_seen = False
    while i < len(lines):
        line = lines[i].rstrip()
        # Tableaux : accumuler les lignes commençant par |
        if line.startswith("|"):
            table_buf.append(line)
            i += 1
            continue
        elif table_buf:
            flush_table(doc, table_buf); table_buf = []

        m = DIAGRAM_RE.match(line.strip())
        if m:
            insert_diagram(doc, m.group(1))
        elif line.startswith("# "):
            if not first_h1_seen:
                first_h1_seen = True  # déjà sur la page de garde
            else:
                doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif re.match(r"^\s*[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, re.sub(r"^\s*[-*]\s+", "", line))
        elif re.match(r"^\s*\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\s*\d+\.\s+", "", line))
        elif line.startswith(">"):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.8)
            add_runs(p, line.lstrip("> ").strip())
        elif line.strip() in ("---", "***", "___"):
            pass
        elif line.strip() == "":
            pass
        else:
            add_runs(doc.add_paragraph(), line)
        i += 1

    if table_buf:
        flush_table(doc, table_buf)

    os.makedirs(DOCX_OUT, exist_ok=True)
    doc.save(docx_path)
    return docx_path


def main():
    mds = sorted(glob.glob(os.path.join(SOURCES, "*.md")))
    if not mds:
        print("Aucun fichier .md dans sources/."); return
    for md in mds:
        name = os.path.splitext(os.path.basename(md))[0]
        out = os.path.join(DOCX_OUT, name + ".docx")
        render(md, out)
        print(f"[OK] {os.path.basename(out)}")
    print(f"\n{len(mds)} document(s) genere(s) dans docx/. Convertissez-les en PDF "
          "(Word/LibreOffice -> Exporter en PDF) puis deposez les PDF dans le repo "
          "Site /assets/docs/.")


if __name__ == "__main__":
    main()
